# -*- coding: utf-8 -*-
import os
import time
from docx import Document
from google import genai
from google.genai import types

# Cole a sua API Key NOVA aqui (dentro das aspas)
API_KEY = "SUA_CHAVE_AQUI"

# Inicializa o cliente com a biblioteca NOVA do Google
client = genai.Client(api_key=API_KEY)

SYSTEM_INSTRUCTIONS = """
Você é um revisor literário profissional e experiente. Sua tarefa é realizar a revisão ortográfica, gramatical, de pontuação e de fluidez textual.

Diretrizes estritas:
1. Mantenha integralmente a voz, o estilo e o enredo do autor.
2. Corrija desvios da norma-padrão da língua portuguesa (crase, regência, concordância, ortografia, pontuação).
3. REGRA DE OURO PARA DIÁLOGOS: Quando o narrador interrompe a fala do personagem (ex: "— Olá — disse ele"), você DEVE aplicar a pontuação correta no trecho do narrador ANTES do segundo travessão para reabrir a fala, respeitando a norma literária de publicação. 
   - Se a fala for a continuação da mesma frase: "— Venha — disse ele —, precisamos ir." (Use vírgula).
   - Se a interrupção dividir frases distintas: "— Venha! — disse ele. — Precisamos ir." (Use ponto final).
4. FLUIDEZ: Se notar palavras repetidas muito próximas de forma não intencional, substitua a segunda ocorrência por um pronome ou sinônimo elegante adequado ao contexto.
5. Ao realizar qualquer correção, você DEVE estruturar o log seguindo EXATAMENTE este formato:
   - [Trecho Original]: "Escreva aqui a frase com o erro" -> [Texto Corrigido]: "Escreva aqui a frase ajustada"

Você deve estruturar a sua resposta exatamente com as duas tags abaixo:

=== TEXTO REVISADO ===
[Insira aqui todo o trecho corrigido, sem introduções]

=== LOG DE ALTERACOES ===
[Insira aqui a lista no formato: - [Trecho Original]: "X" -> [Texto Corrigido]: "Y"]
"""


def ler_documento_em_blocos(caminho_docx, palavras_por_bloco=2500):
    doc = Document(caminho_docx)
    blocos = []
    bloco_atual = []
    contagem_palavras = 0

    for p in doc.paragraphs:
        texto_p = p.text.strip()
        if not texto_p:
            continue

        bloco_atual.append(texto_p)
        contagem_palavras += len(texto_p.split())

        if contagem_palavras >= palavras_por_bloco:
            blocos.append("\n\n".join(bloco_atual))
            bloco_atual = []
            contagem_palavras = 0

    if bloco_atual:
        blocos.append("\n\n".join(bloco_atual))

    return blocos


def processar_revisao(arquivo_entrada, arquivo_texto_revisado, arquivo_log):
    print(f"📖 Lendo o arquivo original: {arquivo_entrada}...")
    try:
        blocos = ler_documento_em_blocos(arquivo_entrada)
    except Exception as e:
        print(f"❌ Erro ao ler o arquivo Word: {e}")
        return

    total_blocos = len(blocos)
    print(
        f"📦 Livro dividido em {total_blocos} blocos de processamento de texto.")

    doc_revisado = Document()
    doc_revisado.add_heading("Romance - Versão Revisada", level=1)

    logs_compilados = []

    print("⚙️  Configurando IA e iniciando sessão de revisão stateless (sem memória)...")

    for i, bloco in enumerate(blocos, start=1):
        print(f"\n🚀 Processando Bloco {i}/{total_blocos}...")

        prompt_envio = f"Por favor, revise o seguinte trecho garantindo fluidez e a pontuação correta dos travessões:\n\n{bloco}"

        tentativas = 3
        sucesso = False
        resposta_texto = ""

        while tentativas > 0 and not sucesso:
            try:
                # CHAMADA STATELESS: O modelo analisa cada bloco do zero
                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=prompt_envio,
                    config=types.GenerateContentConfig(
                        temperature=0.3,
                        top_p=0.95,
                        system_instruction=SYSTEM_INSTRUCTIONS,
                    )
                )
                resposta_texto = response.text
                sucesso = True
            except Exception as e:
                print(
                    f"⚠️  Erro na API: {e}. Tentando novamente em 5 segundos...")
                time.sleep(5)
                tentativas -= 1

        if not sucesso:
            print(
                f"❌ Falha crítica ao processar o Bloco {i} após várias tentativas.")
            continue

        try:
            partes = resposta_texto.split("=== LOG DE ALTERACOES ===")
            texto_revisado_bloco = partes[0].replace(
                "=== TEXTO REVISADO ===", "").strip()
            log_bloco = partes[1].strip() if len(
                partes) > 1 else "- Nenhuma alteração listada ou erro de separação."
        except Exception:
            texto_revisado_bloco = resposta_texto
            log_bloco = f"- Erro ao separar automaticamente o log do bloco {i}."

        for paragrafo in texto_revisado_bloco.split("\n\n"):
            if paragrafo.strip():
                doc_revisado.add_paragraph(paragrafo.strip())

        log_header = f"=== ALTERAÇÕES DO BLOCO {i} ==="
        logs_compilados.append(f"{log_header}\n{log_bloco}\n")

        print(f"✅ Bloco {i} revisado e integrado com sucesso.")
        time.sleep(2)

    try:
        doc_revisado.save(arquivo_texto_revisado)
        print(
            f"\n✨ Livro revisado salvo com sucesso em: {arquivo_texto_revisado}")

        with open(arquivo_log, "w", encoding="utf-8") as f:
            f.write("\n".join(logs_compilados))
        print(f"📝 Relatório completo de alterações salvo em: {arquivo_log}")
        print("\n🎉 Processo concluído! Trabalho pronto para entrega padrão ouro.")
    except Exception as e:
        print(f"❌ Erro ao salvar os arquivos finais: {e}")


if __name__ == "__main__":
    ARQUIVO_ENTRADA = "romance_original.docx"
    ARQUIVO_SAIDA = "romance_revisado_final.docx"
    ARQUIVO_LOG = "log_de_alteracoes_Sarah.txt"

    print("="*60)
    print("   REVISOR AUTOMÁTICO DE OBRAS LITERÁRIAS - GEMINI FLASH")
    print("="*60)

    if not os.path.exists(ARQUIVO_ENTRADA):
        print(f"❌ Arquivo '{ARQUIVO_ENTRADA}' não encontrado.")
    else:
        processar_revisao(ARQUIVO_ENTRADA, ARQUIVO_SAIDA, ARQUIVO_LOG)

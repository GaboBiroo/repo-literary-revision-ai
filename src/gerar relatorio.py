from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def gerar_relatorio_profissional():
    arquivo_txt = "log_de_alteracoes_Sarah.txt"
    arquivo_docx = "Relatorio_de_Revisao_Sarah.docx"

    if not os.path.exists(arquivo_txt):
        print(f"❌ Erro: O arquivo '{arquivo_txt}' não foi encontrado.")
        return

    doc = Document()

    # Capa com estilo limpo
    titulo = doc.add_heading('Relatório de Alterações e Revisão', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.add_run('Manuscrito: Romance da Sarah\n').bold = True
    p_sub.add_run('Revisão Ortográfica e Gramatical\n\n')

    doc.add_page_break()

    doc.add_heading('Lista de Intervenções', level=1)
    doc.add_paragraph('Abaixo, apresento as alterações realizadas no texto, organizadas pelo trecho original encontrado e a respectiva correção aplicada para garantir a norma-padrão da língua portuguesa.')

    with open(arquivo_txt, 'r', encoding='utf-8') as f:
        linhas = f.readlines()

    print("Processando e formatando logs...")

    for linha in linhas:
        linha = linha.strip()
        if not linha or "===" in linha:
            continue

        # O script agora trata especificamente a estrutura [Original] -> [Corrigido]
        if "->" in linha:
            p = doc.add_paragraph(style='List Bullet')
            # Remove o hífen inicial da lista, se houver
            texto_limpo = linha.lstrip('- ').strip()

            # Divide a string no ponto da seta
            partes = texto_limpo.split("->")

            # Formata: Original em negrito, seta normal, corrigido em negrito
            run_orig = p.add_run(partes[0].strip())
            run_orig.bold = True

            p.add_run("  ➔  ")

            run_corr = p.add_run(partes[1].strip())
            run_corr.bold = True
        else:
            # Caso seja apenas uma observação solta
            doc.add_paragraph(linha)

    doc.save(arquivo_docx)
    print(f"✨ Sucesso! O arquivo '{arquivo_docx}' está pronto e profissional.")


if __name__ == "__main__":
    gerar_relatorio_profissional()
